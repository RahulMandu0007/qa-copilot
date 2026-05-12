import streamlit as st
import os
import time
import json
import re
from dotenv import load_dotenv
from anthropic import Anthropic
from docx import Document
from pypdf import PdfReader
from bs4 import BeautifulSoup
import pandas as pd
import chardet
import io

# ================= ENV =================
load_dotenv()
API_KEY = os.getenv("ANTHROPIC_API_KEY")

if not API_KEY:
    st.error("❌ Missing API Key in .env")
    st.stop()

client = Anthropic(api_key=API_KEY)

# ================= AUTH =================
AUTHORIZED_USERS = {
    "rahul": "password123",
    "admin": "admin123"
}

def login():
    st.title("🔐 Secure Login")
    u = st.text_input("Username")
    p = st.text_input("Password", type="password")

    if st.button("Login"):
        if u in AUTHORIZED_USERS and AUTHORIZED_USERS[u] == p:
            st.session_state.auth = True
            st.session_state.user = u
            st.success("✅ Logged in")
            st.rerun()
        else:
            st.error("❌ Invalid credentials")

if "auth" not in st.session_state:
    st.session_state.auth = False

if not st.session_state.auth:
    login()
    st.stop()

# ================= CONFIG =================
MODELS = [
    "claude-haiku-4-5-20251001",
    "claude-sonnet-4-6",
    "claude-opus-4-7"
]

MAX_CHUNK = 2000
MAX_FILE_MB = 5

# ================= SECURITY =================
def mask_sensitive(text):
    text = re.sub(r"\b\d{10}\b", "[PHONE]", text)
    text = re.sub(r"\b[\w\.-]+@[\w\.-]+\b", "[EMAIL]", text)
    text = re.sub(r"\b\d{12}\b", "[ID]", text)
    return text

def log(user, action):
    with open("audit.log", "a") as f:
        f.write(f"{time.ctime()} | {user} | {action}\n")

# ================= RATE LIMIT =================
if "last_call" not in st.session_state:
    st.session_state.last_call = 0

def rate_limit():
    now = time.time()
    if now - st.session_state.last_call < 2:
        st.warning("⏳ Slow down")
        return False
    st.session_state.last_call = now
    return True

# ================= FILE =================
def extract(file):
    name = file.name.lower()

    try:
        if name.endswith(".docx"):
            return "\n".join(p.text for p in Document(file).paragraphs)

        if name.endswith(".pdf"):
            reader = PdfReader(file)
            return "\n".join(p.extract_text() for p in reader.pages if p.extract_text())

        if name.endswith(".csv"):
            return pd.read_csv(file).to_string(index=False)

        if name.endswith((".xlsx", ".xls")):
            sheets = pd.read_excel(file, sheet_name=None)
            return "\n".join(df.to_string(index=False) for df in sheets.values())

        if name.endswith(".html"):
            return BeautifulSoup(file.read(), "html.parser").get_text()

        if name.endswith((".txt", ".md")):
            raw = file.read()
            enc = chardet.detect(raw)["encoding"]
            return raw.decode(enc or "utf-8")

    except:
        return ""

    return ""

def chunk(text):
    return [text[i:i+MAX_CHUNK] for i in range(0, len(text), MAX_CHUNK)]

def retrieve(q, chunks):
    qset = set(q.lower().split())
    scores = [(len(qset & set(c.lower().split())), c) for c in chunks]
    return "\n\n".join(c for s, c in sorted(scores, reverse=True)[:5] if s > 0)

# ================= TABLE =================
def parse_table(text):
    rows = []
    for l in text.split("\n"):
        if "|" in l:
            rows.append([x.strip() for x in l.split("|")])
    if rows:
        return pd.DataFrame(rows[1:], columns=rows[0])
    return None

# ================= AI =================
def complexity(q):
    return 0 if len(q) < 80 else 1 if len(q) < 250 else 2

def tokens(text):
    return int(len(text.split()) * 1.3)

def call_ai(prompt, idx):
    start = time.time()

    for i in range(idx, len(MODELS)):
        try:
            safe_prompt = mask_sensitive(prompt)

            res = client.messages.create(
                model=MODELS[i],
                max_tokens=1500,
                messages=[{"role": "user", "content": safe_prompt}]
            )

            ans = res.content[0].text
            tks = tokens(prompt) + tokens(ans)
            cost = round(tks * 0.000002, 6)

            return ans, MODELS[i], tks, cost, round(time.time()-start, 2)

        except:
            continue

    return "❌ All models failed", "none", 0, 0, 0

# ================= PROMPT =================
def build(action, ctx, hist, q):
    base = f"""
You are a Senior QA expert.

Context:
{ctx}

Conversation:
{hist}
"""
    if action == "test_cases":
        return base + "Generate test cases table."

    if action == "test_data":
        return base + "Generate structured test data."

    if action == "gap":
        return base + "Find missing scenarios."

    if action == "risk":
        return base + "Analyze risks."

    return base + f"\nQuestion:\n{q}"

# ================= UI =================
st.title("🔐 Enterprise AI QA Assistant")

if "chat" not in st.session_state:
    st.session_state.chat = []

if "chunks" not in st.session_state:
    st.session_state.chunks = []

if "answer" not in st.session_state:
    st.session_state.answer = ""

# FILE UPLOAD
files = st.file_uploader("Upload Documents", accept_multiple_files=True)

if files:
    st.session_state.chunks = []
    for f in files:
        if f.size > MAX_FILE_MB * 1024*1024:
            st.error(f"{f.name} too large")
            st.stop()
        st.session_state.chunks += chunk(extract(f))

    log(st.session_state.user, "Uploaded files")

# BUTTONS
col1, col2, col3, col4 = st.columns(4)

action = None
if col1.button("Test Cases"): action="test_cases"
if col2.button("Test Data"): action="test_data"
if col3.button("Gap Analysis"): action="gap"
if col4.button("Risk Analysis"): action="risk"

# CHAT
q = st.chat_input("Ask QA question")

if q or action:
    if not rate_limit():
        st.stop()

    query = q if q else action

    ctx = retrieve(query, st.session_state.chunks) if st.session_state.chunks else ""
    hist = "\n".join(f"{r}:{m}" for r,m in st.session_state.chat[-5:])

    prompt = build(action, ctx, hist, q)
    idx = complexity(query)

    ans, model, tks, cost, t = call_ai(prompt, idx)

    st.session_state.chat.append(("user", query))
    st.session_state.chat.append(("assistant", ans))
    st.session_state.answer = ans

    log(st.session_state.user, query)

    st.info(f"Model: {model} | Tokens: {tks} | Cost: ${cost} | Time: {t}s")

# DISPLAY
for r,m in st.session_state.chat:
    with st.chat_message(r):
        st.write(m)

# EXPORT
if st.session_state.answer:
    df = parse_table(st.session_state.answer)

    c1, c2, c3 = st.columns(3)

    if df is not None:
        b = io.BytesIO()
        df.to_excel(b, index=False)
        c1.download_button("Excel", b.getvalue(), "out.xlsx")
        c2.download_button("CSV", df.to_csv(index=False))

    d = Document()
    d.add_heading("Output")

    if df is not None:
        t = d.add_table(rows=1, cols=len(df.columns))
        for i,c in enumerate(df.columns): t.rows[0].cells[i].text=c
        for _,r in df.iterrows():
            cells = t.add_row().cells
            for i,v in enumerate(r): cells[i].text=str(v)
    else:
        d.add_paragraph(st.session_state.answer)

    buf = io.BytesIO()
    d.save(buf)
    c3.download_button("Word", buf.getvalue(), "out.docx")

# SESSION
st.divider()

if st.button("💾 Save"):
    with open("session.json","w") as f:
        json.dump(st.session_state.chat,f)
    st.success("Saved")

if st.button("📂 Load"):
    try:
        with open("session.json") as f:
            st.session_state.chat = json.load(f)
        st.success("Loaded")
    except:
        st.error("No session")

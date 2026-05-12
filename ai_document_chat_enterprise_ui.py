import streamlit as st
import os
import time
import json
import re
from dotenv import load_dotenv
from anthropic import Anthropic
from docx import Document
from pypdf import PdfReader
from bs4 import BeautifulSoup
import pandas as pd
import chardet
import io

# ================= PAGE CONFIG =================
st.set_page_config(
    page_title="AI QA Assistant",
    page_icon="🤖",
    layout="wide"
)

# ================= ENV =================
load_dotenv()
API_KEY = os.getenv("ANTHROPIC_API_KEY")

if not API_KEY:
    st.error("❌ Missing API Key in .env")
    st.stop()

client = Anthropic(api_key=API_KEY)

# ================= AUTH =================
AUTHORIZED_USERS = {
    "rahul": "password123",
    "admin": "admin123"
}

def login():
    st.title("🔐 Enterprise Login")
    u = st.text_input("Username")
    p = st.text_input("Password", type="password")

    if st.button("Login"):
        if u in AUTHORIZED_USERS and AUTHORIZED_USERS[u] == p:
            st.session_state.auth = True
            st.session_state.user = u
            st.success("✅ Logged in")
            st.rerun()
        else:
            st.error("Invalid credentials")

if "auth" not in st.session_state:
    st.session_state.auth = False

if not st.session_state.auth:
    login()
    st.stop()

# ================= MODELS (YOUR SET) =================
MODELS = {
    "fast": "claude-haiku-4-5-20251001",
    "balanced": "claude-sonnet-4-6",
    "powerful": "claude-opus-4-7"
}

FALLBACK = list(MODELS.values())

# ================= SIDEBAR =================
st.sidebar.title("⚙️ Controls")

mode = st.sidebar.radio(
    "Mode",
    ["Auto", "Force Fast", "Force Balanced", "Force Powerful"]
)

file_limit = st.sidebar.slider("Max File Size (MB)", 1, 10, 5)

st.sidebar.markdown("---")
st.sidebar.info("💡 Auto mode selects best model based on query complexity")

# ================= SECURITY =================
def mask_sensitive(text):
    text = re.sub(r"\b\d{10}\b", "[PHONE]", text)
    text = re.sub(r"\b[\w\.-]+@[\w\.-]+\b", "[EMAIL]", text)
    text = re.sub(r"\b\d{12}\b", "[ID]", text)
    return text

def log(user, action):
    with open("audit.log", "a") as f:
        f.write(f"{time.ctime()} | {user} | {action}\n")

# ================= RATE LIMIT =================
if "last_call" not in st.session_state:
    st.session_state.last_call = 0

def rate_limit():
    now = time.time()
    if now - st.session_state.last_call < 2:
        st.warning("⏳ Please wait...")
        return False
    st.session_state.last_call = now
    return True

# ================= FILE HANDLING =================
def extract(file):
    name = file.name.lower()
    try:
        if name.endswith(".docx"):
            return "\n".join(p.text for p in Document(file).paragraphs)

        if name.endswith(".pdf"):
            reader = PdfReader(file)
            return "\n".join(p.extract_text() for p in reader.pages if p.extract_text())

        if name.endswith(".csv"):
            return pd.read_csv(file).to_string(index=False)

        if name.endswith((".xlsx", ".xls")):
            sheets = pd.read_excel(file, sheet_name=None)
            return "\n".join(df.to_string(index=False) for df in sheets.values())

        if name.endswith(".html"):
            return BeautifulSoup(file.read(), "html.parser").get_text()

        if name.endswith((".txt", ".md")):
            raw = file.read()
            enc = chardet.detect(raw)["encoding"]
            return raw.decode(enc or "utf-8")

    except:
        return ""

    return ""

def chunk(text):
    return [text[i:i+2000] for i in range(0, len(text), 2000)]

def retrieve(q, chunks):
    qset = set(q.lower().split())
    scores = [(len(qset & set(c.lower().split())), c) for c in chunks]
    return "\n\n".join(c for s, c in sorted(scores, reverse=True)[:5] if s > 0)

# ================= TABLE =================
def parse_table(text):
    rows = []
    for l in text.split("\n"):
        if "|" in l:
            rows.append([x.strip() for x in l.split("|")])
    if rows:
        return pd.DataFrame(rows[1:], columns=rows[0])
    return None

# ================= MODEL SELECT =================
def choose_model(query, action, ctx):
    if mode == "Force Fast":
        return MODELS["fast"]
    if mode == "Force Balanced":
        return MODELS["balanced"]
    if mode == "Force Powerful":
        return MODELS["powerful"]

    # AUTO MODE
    if len(ctx) > 4000:
        return MODELS["powerful"]

    if len(query) > 300:
        return MODELS["powerful"]

    if action in ["test_cases", "test_data", "gap", "risk"]:
        return MODELS["balanced"]

    if len(query) < 80:
        return MODELS["fast"]

    return MODELS["balanced"]

# ================= AI =================
def tokens(text):
    return int(len(text.split()) * 1.3)

def call_ai(prompt, model):
    start = time.time()
    tried = set()

    for m in [model] + FALLBACK:
        if m in tried:
            continue
        tried.add(m)

        try:
            safe_prompt = mask_sensitive(prompt)

            res = client.messages.create(
                model=m,
                max_tokens=1500,
                messages=[{"role": "user", "content": safe_prompt}]
            )

            ans = res.content[0].text
            tks = tokens(prompt) + tokens(ans)
            cost = round(tks * 0.000002, 6)

            return ans, m, tks, cost, round(time.time()-start, 2)

        except:
            continue

    return "❌ All models failed", "none", 0, 0, 0

# ================= PROMPT =================
def build(action, ctx, hist, q):
    base = f"""
You are a Senior QA expert.

Context:
{ctx}

Conversation:
{hist}
"""

    if action == "test_cases":
        return base + "Generate detailed test cases in table format."

    if action == "test_data":
        return base + "Generate structured test data."

    if action == "gap":
        return base + "Identify missing scenarios."

    if action == "risk":
        return base + "Analyze risks."

    return base + f"\nQuestion:\n{q}"

# ================= MAIN UI =================
st.title("🚀 Enterprise AI QA Assistant")

if "chat" not in st.session_state:
    st.session_state.chat = []

if "chunks" not in st.session_state:
    st.session_state.chunks = []

if "answer" not in st.session_state:
    st.session_state.answer = ""

# FILE UPLOAD
files = st.file_uploader("📂 Upload Documents", accept_multiple_files=True)

if files:
    st.session_state.chunks = []
    for f in files:
        if f.size > file_limit * 1024 * 1024:
            st.error(f"{f.name} too large")
            st.stop()
        st.session_state.chunks += chunk(extract(f))

    log(st.session_state.user, "Files uploaded")

# ACTION BUTTONS
st.subheader("⚡ QA Actions")

col1, col2, col3, col4 = st.columns(4)

action = None
if col1.button("✅ Test Cases"): action="test_cases"
if col2.button("📊 Test Data"): action="test_data"
if col3.button("🔍 Gap Analysis"): action="gap"
if col4.button("⚠️ Risk Analysis"): action="risk"

# CHAT INPUT
q = st.chat_input("Ask your QA question...")

if q or action:
    if not rate_limit():
        st.stop()

    query = q if q else action
    ctx = retrieve(query, st.session_state.chunks) if st.session_state.chunks else ""
    hist = "\n".join(f"{r}:{m}" for r, m in st.session_state.chat[-5:])

    prompt = build(action, ctx, hist, q)
    model = choose_model(query, action, ctx)

    ans, used_model, tks, cost, t = call_ai(prompt, model)

    st.session_state.chat.append(("user", query))
    st.session_state.chat.append(("assistant", ans))
    st.session_state.answer = ans

    log(st.session_state.user, query)

    # METRICS DISPLAY
    st.sidebar.markdown("### 📊 Metrics")
    st.sidebar.write(f"Model: {used_model}")
    st.sidebar.write(f"Tokens: {tks}")
    st.sidebar.write(f"Cost: ${cost}")
    st.sidebar.write(f"Time: {t}s")

# CHAT DISPLAY
for r, m in st.session_state.chat:
    with st.chat_message(r):
        st.write(m)

# EXPORT
if st.session_state.answer:
    st.subheader("⬇️ Export")

    df = parse_table(st.session_state.answer)

    col1, col2, col3 = st.columns(3)

    if df is not None:
        buffer = io.BytesIO()
        df.to_excel(buffer, index=False)
        col1.download_button("Excel", buffer.getvalue(), "output.xlsx")

        col2.download_button("CSV", df.to_csv(index=False), "output.csv")

    doc = Document()
    doc.add_heading("AI Output")

    if df is not None:
        t = doc.add_table(rows=1, cols=len(df.columns))
        for i, c in enumerate(df.columns):
            t.rows[0].cells[i].text = c
        for _, row in df.iterrows():
            cells = t.add_row().cells
            for i, val in enumerate(row):
                cells[i].text = str(val)
    else:
        doc.add_paragraph(st.session_state.answer)

    mem = io.BytesIO()
    doc.save(mem)
    col3.download_button("Word", mem.getvalue(), "output.docx")

# SESSION
st.divider()

colA, colB = st.columns(2)

if colA.button("💾 Save Session"):
    with open("session.json", "w") as f:
        json.dump(st.session_state.chat, f)
    st.success("Saved")

if colB.button("📂 Load Session"):
    try:
        with open("session.json") as f:
            st.session_state.chat = json.load(f)
        st.success("Loaded")
    except:
        st.error("No session found")

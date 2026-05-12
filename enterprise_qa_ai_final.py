import streamlit as st
import os, time, json, re, io
import pandas as pd
from dotenv import load_dotenv
from anthropic import Anthropic
from docx import Document
from pypdf import PdfReader
from bs4 import BeautifulSoup
import chardet

# ========== CONFIG ==========
st.set_page_config(page_title="Enterprise QA AI", layout="wide")

st.markdown("""
<style>
.main-title {font-size:28px;font-weight:700;color:#2c3e50;}
.sub-title {color:gray;margin-bottom:15px;}
</style>
""", unsafe_allow_html=True)

st.markdown('<div class="main-title">🚀 Enterprise QA AI Assistant</div>', unsafe_allow_html=True)
st.markdown('<div class="sub-title">End-to-End QA Automation System</div>', unsafe_allow_html=True)

# ========== ENV ==========
load_dotenv()
API_KEY = os.getenv("ANTHROPIC_API_KEY")
if not API_KEY:
    st.error("Missing API key")
    st.stop()

client = Anthropic(api_key=API_KEY)

# ========== AUTH ==========
USERS = {"rahul": "password123"}

if "auth" not in st.session_state:
    st.session_state.auth = False

if not st.session_state.auth:
    st.header("🔐 Login")
    u = st.text_input("Username")
    p = st.text_input("Password", type="password")

    if st.button("Login"):
        if USERS.get(u) == p:
            st.session_state.auth = True
            st.rerun()
        else:
            st.error("Invalid credentials")
    st.stop()

# ========== MODELS ==========
MODELS = {
    "fast": "claude-haiku-4-5-20251001",
    "balanced": "claude-sonnet-4-6",
    "powerful": "claude-opus-4-7"
}

# ========== SIDEBAR ==========
st.sidebar.title("⚙️ Controls")

mode = st.sidebar.selectbox(
    "Model Mode",
    ["Auto", "Fast", "Balanced", "Powerful"]
)

# ========== SECURITY ==========
def mask(text):
    text = re.sub(r"\S+@\S+", "[EMAIL]", text)
    text = re.sub(r"\b\d{10}\b", "[PHONE]", text)
    return text

# ========== FILE ==========
def extract(file):
    name = file.name.lower()
    try:
        if name.endswith(".docx"):
            return "\n".join(p.text for p in Document(file).paragraphs)
        if name.endswith(".pdf"):
            return "\n".join(p.extract_text() for p in PdfReader(file).pages if p.extract_text())
        if name.endswith(".csv"):
            return pd.read_csv(file).to_string()
        if name.endswith(".html"):
            return BeautifulSoup(file.read(),"html.parser").get_text()
        if name.endswith((".txt",".md")):
            raw=file.read()
            enc=chardet.detect(raw)["encoding"]
            return raw.decode(enc or "utf-8")
    except: return ""
    return ""

def chunk(text):
    return [text[i:i+2000] for i in range(0,len(text),2000)]

def retrieve(q,chunks):
    qset=set(q.lower().split())
    scored=[(len(qset & set(c.lower().split())),c) for c in chunks]
    return "\n\n".join(c for s,c in sorted(scored,reverse=True)[:5])

# ========== MODEL SELECT ==========
def choose_model(q,ctx):
    if mode!="Auto":
        return MODELS[mode.lower()]
    if len(q)>300 or len(ctx)>4000:
        return MODELS["powerful"]
    if len(q)<80:
        return MODELS["fast"]
    return MODELS["balanced"]

# ========== AI ==========
def call_ai(prompt,model):
    try:
        res=client.messages.create(
            model=model,
            max_tokens=2000,
            messages=[{"role":"user","content":mask(prompt)}]
        )
        return res.content[0].text
    except:
        return "Error"

# ========== PROMPT REFINEMENT ==========
def refine_prompt(user_prompt):
    return f"""
You are a Senior QA Analyst.
Improve and structure the following prompt before answering:

{user_prompt}
"""

# ========== TEST CASE FORMAT ==========
def build_prompt(context):
    return f"""
Generate structured QA test cases in JSON format:

[Test ID, Story, Test Title, Pre-conditions, Steps,
Expected Result, Test Data, Priority, Test Type,
Automation, Status]

Context:
{context}

Return ONLY JSON array.
"""

def parse_json(text):
    try:
        js=text[text.find("["):text.rfind("]")+1]
        return pd.DataFrame(json.loads(js))
    except:
        return None

# ========== EXCEL BUILDER ==========
def build_excel(df):
    output=io.BytesIO()
    with pd.ExcelWriter(output, engine='xlsxwriter') as writer:
        workbook=writer.book

        header_fmt=workbook.add_format({'bold':True,'bg_color':'#2c3e50','font_color':'white'})
        high=workbook.add_format({'bg_color':'#ff6b6b'})
        med=workbook.add_format({'bg_color':'#ffd93d'})
        low=workbook.add_format({'bg_color':'#6bcb77'})

        # Summary
        summary=pd.DataFrame({
            "Metric":["Total","High","Medium","Low"],
            "Value":[len(df),"","",""]
        })
        summary.to_excel(writer,"Cover",index=False)

        # Coverage
        cov=df.groupby("Story")["Test ID"].count().reset_index()
        cov.columns=["Story","Count"]
        cov.to_excel(writer,"Coverage",index=False)

        types={
            "01 Functional Tests":"Functional",
            "02 API Tests":"API",
            "03 DB Tests":"DB",
            "04 Integration Tests":"Integration"
        }

        for name,t in types.items():
            sub=df[df["Test Type"].str.contains(t,case=False,na=False)]
            if sub.empty: continue

            sub.to_excel(writer,name,index=False)
            ws=writer.sheets[name]

            for col_num,val in enumerate(sub.columns):
                ws.write(0,col_num,val,header_fmt)

            ws.autofilter(0,0,len(sub),len(sub.columns)-1)
            ws.freeze_panes(1,0)

            if "Priority" in sub.columns:
                col=sub.columns.get_loc("Priority")
                ws.conditional_format(1,col,len(sub),col,{'type':'text','criteria':'containing','value':'High','format':high})
                ws.conditional_format(1,col,len(sub),col,{'type':'text','criteria':'containing','value':'Medium','format':med})
                ws.conditional_format(1,col,len(sub),col,{'type':'text','criteria':'containing','value':'Low','format':low})

    return output

# ========== STATE ==========
if "chunks" not in st.session_state: st.session_state.chunks=[]
if "answer" not in st.session_state: st.session_state.answer=""

# ========== FILE UPLOAD ==========
st.subheader("📂 Upload Documents")
files=st.file_uploader("",accept_multiple_files=True)

if files:
    st.session_state.chunks=[]
    for f in files:
        if f.size>5*1024*1024:
            st.error("File too large")
            st.stop()
        st.session_state.chunks+=chunk(extract(f))

# ========== ACTIONS ==========
st.subheader("⚡ Actions")

col1,col2=st.columns(2)

if col1.button("✅ Generate Test Pack"):
    ctx=retrieve("test cases",st.session_state.chunks)
    model=choose_model("test cases",ctx)
    ans=call_ai(build_prompt(ctx),model)
    st.session_state.answer=ans

# ========== CHAT ==========
st.subheader("💬 Chat")

user_q=st.text_input("Ask anything...")

if st.button("Send"):
    refined=refine_prompt(user_q)
    model=choose_model(user_q,"")
    ans=call_ai(refined,model)
    st.session_state.answer=ans

# ========== OUTPUT ==========
if st.session_state.answer:
    st.markdown("### 🧠 AI Response")
    st.write(st.session_state.answer)

    df=parse_json(st.session_state.answer)

    if df is not None:
        excel=build_excel(df)
        st.download_button("📥 Download Excel Test Pack",excel.getvalue(),"QA_TestPack.xlsx")
    else:
        doc=Document()
        doc.add_paragraph(st.session_state.answer)
        buf=io.BytesIO()
        doc.save(buf)
        st.download_button("📥 Download Output (Word)",buf.getvalue(),"output.docx")
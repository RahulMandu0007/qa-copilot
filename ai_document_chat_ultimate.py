import streamlit as st
import os
import time
import json
from dotenv import load_dotenv
from anthropic import Anthropic
from docx import Document
from pypdf import PdfReader
from bs4 import BeautifulSoup
import pandas as pd
import chardet
import io

# ---------------- ENV ----------------
load_dotenv()
API_KEY = os.getenv("ANTHROPIC_API_KEY")

if not API_KEY:
    st.error("❌ API key missing in .env file")
    st.stop()

client = Anthropic(api_key=API_KEY)

# ---------------- CONFIG ----------------
MODELS = [
    "claude-haiku-4-5-20251001",
    "claude-sonnet-4-6",
    "claude-opus-4-7"
]
MAX_CHUNK_SIZE = 2000

# ---------------- FILE EXTRACTION ----------------
def extract_text(file):
    name = file.name.lower()
    try:
        if name.endswith((".txt", ".md", ".log")):
            raw = file.read()
            enc = chardet.detect(raw)["encoding"]
            return raw.decode(enc or "utf-8", errors="ignore")

        elif name.endswith(".docx"):
            doc = Document(file)
            return "\n".join(p.text for p in doc.paragraphs)

        elif name.endswith(".pdf"):
            reader = PdfReader(file)
            return "\n".join(p.extract_text() for p in reader.pages if p.extract_text())

        elif name.endswith((".xlsx", ".xls")):
            sheets = pd.read_excel(file, sheet_name=None)
            return "\n".join(df.to_string(index=False) for df in sheets.values())

        elif name.endswith(".csv"):
            return pd.read_csv(file).to_string(index=False)

        elif name.endswith(".html"):
            return BeautifulSoup(file.read(), "html.parser").get_text()

        elif name.endswith((".json", ".xml")):
            return file.read().decode("utf-8", errors="ignore")

    except:
        return ""

    return ""

# ---------------- CHUNK ----------------
def chunk(text):
    return [text[i:i+MAX_CHUNK_SIZE] for i in range(0, len(text), MAX_CHUNK_SIZE)]

# ---------------- RETRIEVAL ----------------
def get_relevant_chunks(query, chunks):
    scores = []
    q_words = set(query.lower().split())

    for c in chunks:
        c_words = set(c.lower().split())
        score = len(q_words.intersection(c_words))
        scores.append((score, c))

    top = sorted(scores, reverse=True)[:5]
    return "\n\n".join(t[1] for t in top if t[0] > 0)

# ---------------- TABLE DETECT ----------------
def parse_table(text):
    rows = []
    for line in text.split("\n"):
        if "|" in line and "---" not in line:
            cols = [c.strip() for c in line.split("|")]
            if len(cols) > 2:
                rows.append(cols)
    if rows:
        return pd.DataFrame(rows[1:], columns=rows[0])
    return None

# ---------------- COMPLEXITY ----------------
def estimate_complexity(question):
    if len(question) < 80:
        return 0
    elif len(question) < 250:
        return 1
    return 2

def estimate_tokens(text):
    return int(len(text.split()) * 1.3)

# ---------------- API CALL ----------------
def call_claude(prompt, start_model):
    start_time = time.time()

    for i in range(start_model, len(MODELS)):
        try:
            response = client.messages.create(
                model=MODELS[i],
                max_tokens=1500,
                messages=[{"role": "user", "content": prompt}]
            )

            end = time.time()
            answer = response.content[0].text

            tokens = estimate_tokens(prompt) + estimate_tokens(answer)
            cost = round(tokens * 0.000002, 6)

            return answer, MODELS[i], tokens, cost, round(end-start_time, 2)

        except Exception as e:
            continue

    return "❌ All models failed", "none", 0, 0, 0

# ---------------- PROMPT BUILDER ----------------
def build_prompt(action, context, history, question):

    base = f"""
You are a Senior QA expert.

Context:
{context}

Conversation:
{history}
"""

    if action == "test_cases":
        return base + """
Generate detailed manual test cases:

Test Case ID | Title | Preconditions | Steps | Expected Result
"""

    elif action == "test_data":
        return base + "Generate structured test data with edge cases."

    elif action == "gap":
        return base + "Find missing scenarios."

    elif action == "risk":
        return base + "Perform risk analysis."

    else:
        return base + f"\nQuestion:\n{question}"

# ---------------- UI ----------------
st.set_page_config(layout="wide")
st.title("🚀 AI QA Assistant – Ultimate")

# State
if "chunks" not in st.session_state:
    st.session_state.chunks = []

if "chat" not in st.session_state:
    st.session_state.chat = []

if "last_answer" not in st.session_state:
    st.session_state.last_answer = ""

# File upload
files = st.file_uploader("Upload Documents", accept_multiple_files=True)

if files:
    st.session_state.chunks = []
    for f in files:
        st.session_state.chunks += chunk(extract_text(f))
    st.success("✅ Documents loaded")

# QA Buttons
st.subheader("⚡ Quick QA Actions")
colA, colB, colC, colD = st.columns(4)

action = None
if colA.button("Generate Test Cases"): action="test_cases"
if colB.button("Generate Test Data"): action="test_data"
if colC.button("Missing Scenarios"): action="gap"
if colD.button("Risk Analysis"): action="risk"

# Chat
question = st.chat_input("Ask anything...")

if question or action:
    query = question if question else action
    st.session_state.chat.append(("user", query))

    context = ""
    if st.session_state.chunks:
        context = get_relevant_chunks(query, st.session_state.chunks)

    history = "\n".join(f"{r}: {m}" for r, m in st.session_state.chat[-5:])

    prompt = build_prompt(action, context, history, question)

    complexity = estimate_complexity(query)

    answer, model, tokens, cost, t = call_claude(prompt, complexity)

    st.session_state.chat.append(("assistant", answer))
    st.session_state.last_answer = answer

    st.info(f"Model: {model} | Tokens: {tokens} | Cost: ${cost} | Time: {t}s")

# Display
for role, msg in st.session_state.chat:
    with st.chat_message(role):
        st.markdown(msg)

# Export
if st.session_state.last_answer:
    st.subheader("⬇️ Export")
    text = st.session_state.last_answer
    df = parse_table(text)

    col1, col2, col3 = st.columns(3)

    if df is not None:
        buf = io.BytesIO()
        df.to_excel(buf, index=False)
        col1.download_button("Excel", buf.getvalue(), "out.xlsx")

        col2.download_button("CSV", df.to_csv(index=False), "out.csv")

    doc = io.BytesIO()
    d = Document()
    d.add_heading("Output")
    if df is not None:
        t = d.add_table(rows=1, cols=len(df.columns))
        for i,c in enumerate(df.columns): t.rows[0].cells[i].text=c
        for _,r in df.iterrows():
            cells = t.add_row().cells
            for i,v in enumerate(r): cells[i].text=str(v)
    else:
        d.add_paragraph(text)
    d.save(doc)
    col3.download_button("Word", doc.getvalue(), "out.docx")

# Session Save/Load
st.divider()
if st.button("💾 Save"):
    with open("session.json","w") as f:
        json.dump(st.session_state.chat,f)
    st.success("Saved")

if st.button("📂 Load"):
    try:
        with open("session.json") as f:
            st.session_state.chat = json.load(f)
        st.success("Loaded")
    except:
        st.error("No session")